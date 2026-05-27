#!/usr/bin/env python3
"""
Konverton DOKUMENTIMI.md në DOKUMENTIMI.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD = open("DOKUMENTIMI.md", encoding="utf-8").read()
doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# --- Styles ---
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def set_heading_style(para, level):
    colors = {1: '1e3a5f', 2: 'c45911', 3: '2e75b6', 4: '375623'}
    sizes  = {1: 20, 2: 16, 3: 13, 4: 12}
    para.style = doc.styles[f'Heading {level}']
    run = para.runs[0] if para.runs else para.add_run(para.text)
    run.font.color.rgb = RGBColor.from_string(colors.get(level, '000000'))
    run.font.size = Pt(sizes.get(level, 12))
    run.font.bold = True

def add_code_block(text):
    """Shto bllok kodi me sfond gri."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    # Sfond gri nëpërmjet shading XML
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1f, 0x2d, 0x3d)
    return para

def add_table_from_md(lines):
    """Parse tabelë Markdown dhe shto në dokument."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or set(line.replace('|','').replace('-','').replace(' ','')) == set():
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= ncols:
                break
            cell = row.cells[j]
            # Pastro markdown bold/code nga teksti i qelizës
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            clean = re.sub(r'`(.+?)`',       r'\1', clean)
            cell.text = clean
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if i == 0:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                # Header sfond e kaltër e errët
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  '1e3a5f')
                tcPr.append(shd)
    doc.add_paragraph()

def inline_format(para, text):
    """Apliko bold/italic/code brenda një paragrafi."""
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            if part:
                para.add_run(part)

# ----------------------------------------------------------------
# Parse Markdown rresht pas rreshti
# ----------------------------------------------------------------
lines = MD.split('\n')
i = 0
in_code = False
code_buf = []
in_table = False
table_buf = []

while i < len(lines):
    line = lines[i]

    # --- Bllok kodi ---
    if line.strip().startswith('```'):
        if not in_code:
            in_code = True
            code_buf = []
        else:
            add_code_block('\n'.join(code_buf))
            in_code = False
            code_buf = []
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # --- Tabelë ---
    if line.strip().startswith('|'):
        table_buf.append(line)
        i += 1
        continue
    else:
        if table_buf:
            add_table_from_md(table_buf)
            table_buf = []

    # --- Tituj ---
    if line.startswith('#### '):
        para = doc.add_paragraph(line[5:].strip())
        set_heading_style(para, 4)
        i += 1
        continue

    if line.startswith('### '):
        para = doc.add_paragraph(line[4:].strip())
        set_heading_style(para, 3)
        i += 1
        continue

    if line.startswith('## '):
        doc.add_paragraph()
        para = doc.add_paragraph(line[3:].strip())
        set_heading_style(para, 2)
        i += 1
        continue

    if line.startswith('# '):
        para = doc.add_paragraph(line[2:].strip())
        set_heading_style(para, 1)
        i += 1
        continue

    # --- Ndarëse horizontale ---
    if line.strip() == '---':
        para = doc.add_paragraph()
        pPr  = para._p.get_or_add_pPr()
        pb   = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'CCCCCC')
        pb.append(bottom)
        pPr.append(pb)
        i += 1
        continue

    # --- Lista me pika ---
    if re.match(r'^[\-\*] ', line):
        text = line[2:].strip()
        para = doc.add_paragraph(style='List Bullet')
        inline_format(para, text)
        i += 1
        continue

    # --- Lista e numëruar ---
    if re.match(r'^\d+\. ', line):
        text = re.sub(r'^\d+\. ', '', line).strip()
        para = doc.add_paragraph(style='List Number')
        inline_format(para, text)
        i += 1
        continue

    # --- Rresht bosh ---
    if line.strip() == '':
        i += 1
        continue

    # --- Paragraf normal ---
    if line.strip():
        para = doc.add_paragraph()
        inline_format(para, line.strip())

    i += 1

# Mbaro tabelën nëse mbeti e hapur
if table_buf:
    add_table_from_md(table_buf)

doc.save('DOKUMENTIMI.docx')
print("✓ DOKUMENTIMI.docx u krijua me sukses!")
